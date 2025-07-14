#!/usr/bin/env python3
"""
CSRR Faculty Tracker - Rutgers Law Themed Dashboard
Real automation system for tracking faculty publications
"""

from flask import Flask, render_template_string, request, redirect, url_for, flash, jsonify, send_file
import pandas as pd
from datetime import datetime, timedelta
import os
import sys
import json
from pathlib import Path
import threading
import time

# Add the parent directory to path
sys.path.append('/Users/azrabano')
from csrr_faculty_tracker import CSRRFacultyTracker

app = Flask(__name__)
app.config['SECRET_KEY'] = 'csrr-tracker-secret-key'

# Initialize tracker
tracker = CSRRFacultyTracker()

# In-memory storage for demo (replace with database in production)
search_history = []
email_subscribers = []

@app.route('/')
def dashboard():
    """Rutgers Law themed dashboard with real data"""
    
    # Real analytics data only
    analytics = {
        'total_faculty': len(tracker.faculty_names),
        'total_searches': len(search_history),
        'total_subscribers': len(email_subscribers),
        'last_search': search_history[-1]['date'] if search_history else 'Never'
    }
    
    return render_template_string(DASHBOARD_HTML, analytics=analytics, search_history=search_history[:5])

@app.route('/run-search', methods=['POST'])
def run_search():
    """Start a new search and generate reports - REAL SEARCH"""
    
    # Add to search history
    search_record = {
        'id': len(search_history) + 1,
        'date': datetime.now().strftime('%Y-%m-%d %H:%M'),
        'status': 'Running',
        'results': 0
    }
    search_history.append(search_record)
    
    # Start background search with REAL faculty tracker
    thread = threading.Thread(target=background_search, args=(search_record,))
    thread.daemon = True
    thread.start()
    
    flash('Faculty search started! This will search for publications by all 70+ CSRR faculty affiliates over the last 30 days.', 'success')
    return redirect(url_for('dashboard'))

def background_search(search_record):
    """Run the ACTUAL search in background"""
    try:
        # Run the real search using the faculty tracker
        results_count = tracker.run_monthly_search()
        
        # Update search record with real results
        search_record['status'] = 'Completed'
        search_record['results'] = results_count
        search_record['excel_report'] = f"CSRR_Report_{datetime.now().strftime('%Y%m%d')}.xlsx"
        search_record['word_report'] = f"CSRR_Report_{datetime.now().strftime('%Y%m%d')}.docx"
        
        # Generate Word document in the format you specified
        generate_word_report(search_record, results_count)
        
    except Exception as e:
        search_record['status'] = 'Failed'
        search_record['error'] = str(e)
        print(f"Search failed: {e}")

def generate_word_report(search_record, results_count):
    """Generate Word document in CSRR format for website review"""
    try:
        from docx import Document
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # Title with CSRR branding
        title = doc.add_heading(f'CSRR Faculty Affiliates Publications - {datetime.now().strftime("%B %Y")}', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Header info
        doc.add_paragraph(f'Center for Security, Race and Rights')
        doc.add_paragraph(f'Rutgers Law School')
        doc.add_paragraph(f'Report Generated: {datetime.now().strftime("%B %d, %Y")}')
        doc.add_paragraph('')
        
        # Summary section
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(f'Total Publications Found: {results_count}')
        doc.add_paragraph(f'Search Period: {(datetime.now() - timedelta(days=30)).strftime("%B %d")} - {datetime.now().strftime("%B %d, %Y")}')
        doc.add_paragraph(f'Faculty Affiliates Monitored: {len(tracker.faculty_names)}')
        doc.add_paragraph('')
        
        # Content types searched
        doc.add_heading('Content Types Monitored', level=1)
        doc.add_paragraph('• Op-Eds and Opinion Pieces')
        doc.add_paragraph('• Print Media Interviews')
        doc.add_paragraph('• Television and Radio Interviews')
        doc.add_paragraph('• Podcast Appearances')
        doc.add_paragraph('• Academic Articles and Commentary')
        doc.add_paragraph('')
        
        # Note about real data
        if results_count > 0:
            doc.add_heading('Publications Found', level=1)
            doc.add_paragraph('Detailed publication data has been generated and saved in the Excel report. Please review the Excel file for complete information including:')
            doc.add_paragraph('• Faculty member names and affiliations')
            doc.add_paragraph('• Publication titles and descriptions')
            doc.add_paragraph('• Source publications and URLs')
            doc.add_paragraph('• Publication dates and search terms used')
        else:
            doc.add_heading('No Publications Found', level=1)
            doc.add_paragraph('No new publications were found during this search period. This could be due to:')
            doc.add_paragraph('• Limited recent activity by faculty affiliates')
            doc.add_paragraph('• Search engine limitations or rate limiting')
            doc.add_paragraph('• Changes in publication patterns or timing')
        
        doc.add_paragraph('')
        
        # Recommendations for CSRR in the News
        doc.add_heading('Next Steps for CSRR in the News', level=1)
        doc.add_paragraph('1. Review the accompanying Excel report for detailed publication information')
        doc.add_paragraph('2. Select high-impact publications for featuring on the CSRR website')
        doc.add_paragraph('3. Prioritize publications from major media outlets (Washington Post, NYT, CNN, etc.)')
        doc.add_paragraph('4. Consider multimedia content (TV/radio interviews) for diverse content types')
        doc.add_paragraph('5. Update the CSRR in the News section: https://csrr.rutgers.edu/newsroom/csrr-in-the-news/')
        
        # Save document
        reports_dir = Path('/Users/azrabano/CSRR_Reports')
        reports_dir.mkdir(exist_ok=True)
        doc_path = reports_dir / search_record['word_report']
        doc.save(doc_path)
        
        search_record['word_path'] = str(doc_path)
        
    except ImportError:
        print("python-docx not installed. Installing...")
        os.system("pip install python-docx")
        generate_word_report(search_record, results_count)  # Retry
    except Exception as e:
        print(f"Error generating Word document: {e}")

@app.route('/subscribe', methods=['POST'])
def subscribe():
    """Subscribe to email notifications"""
    email = request.form.get('email')
    if email and email not in email_subscribers:
        email_subscribers.append(email)
        flash(f'Successfully subscribed {email} to monthly CSRR faculty publication reports!', 'success')
    else:
        flash('Email already subscribed or invalid.', 'warning')
    return redirect(url_for('dashboard'))

@app.route('/download-report/<int:search_id>/<report_type>')
def download_report(search_id, report_type):
    """Download Excel or Word report"""
    if search_id <= len(search_history):
        search = search_history[search_id - 1]
        if report_type == 'excel':
            # Find the most recent Excel report
            reports_dir = Path('/Users/azrabano/CSRR_Reports')
            if reports_dir.exists():
                excel_files = list(reports_dir.glob('*.xlsx'))
                if excel_files:
                    latest_excel = max(excel_files, key=lambda f: f.stat().st_mtime)
                    return send_file(latest_excel, as_attachment=True)
            flash('Excel report not found', 'error')
        elif report_type == 'word' and 'word_path' in search:
            return send_file(search['word_path'], as_attachment=True)
    
    flash('Report not found', 'error')
    return redirect(url_for('dashboard'))

@app.route('/faculty-list')
def faculty_list():
    """Display all CSRR faculty members"""
    return render_template_string(FACULTY_LIST_HTML, faculty_names=tracker.faculty_names)

# Rutgers Law themed HTML Templates
DASHBOARD_HTML = '''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>CSRR Faculty Tracker | Rutgers Law</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/css/bootstrap.min.css" rel="stylesheet">
    <link href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0/css/all.min.css" rel="stylesheet">
    <style>
        :root {
            --rutgers-red: #CC0033;
            --rutgers-light-red: #E74C3C;
            --rutgers-dark-red: #A00025;
            --rutgers-gray: #666666;
            --rutgers-light-gray: #F5F5F5;
        }
        
        body { 
            background-color: var(--rutgers-light-gray);
            font-family: 'Arial', sans-serif;
        }
        
        .navbar { 
            background: linear-gradient(135deg, var(--rutgers-red) 0%, var(--rutgers-dark-red) 100%);
            box-shadow: 0 2px 10px rgba(204, 0, 51, 0.3);
        }
        
        .navbar-brand {
            font-weight: bold;
            font-size: 1.5rem;
        }
        
        .card { 
            border: none; 
            border-radius: 15px; 
            box-shadow: 0 4px 20px rgba(0,0,0,0.1); 
            margin-bottom: 20px;
            transition: transform 0.2s ease;
        }
        
        .card:hover {
            transform: translateY(-2px);
        }
        
        .stat-card { 
            background: linear-gradient(135deg, var(--rutgers-red) 0%, var(--rutgers-light-red) 100%); 
            color: white; 
        }
        
        .btn-primary { 
            background: linear-gradient(135deg, var(--rutgers-red) 0%, var(--rutgers-light-red) 100%); 
            border: none;
            transition: all 0.3s ease;
        }
        
        .btn-primary:hover {
            background: linear-gradient(135deg, var(--rutgers-dark-red) 0%, var(--rutgers-red) 100%);
            transform: translateY(-1px);
        }
        
        .search-section { 
            background: linear-gradient(135deg, var(--rutgers-red) 0%, var(--rutgers-light-red) 100%); 
            color: white; 
            border-radius: 20px; 
            padding: 40px; 
            margin-bottom: 30px;
            text-align: center;
        }
        
        .rutgers-logo {
            height: 40px;
            margin-right: 10px;
        }
        
        .feature-icon {
            color: var(--rutgers-red);
        }
        
        .badge.bg-success { background-color: #28a745 !important; }
        .badge.bg-primary { background-color: var(--rutgers-red) !important; }
        .badge.bg-danger { background-color: #dc3545 !important; }
        
        .table-hover tbody tr:hover {
            background-color: rgba(204, 0, 51, 0.05);
        }
        
        .alert-success {
            background-color: rgba(40, 167, 69, 0.1);
            border-color: #28a745;
            color: #155724;
        }
    </style>
</head>
<body>
    <nav class="navbar navbar-dark">
        <div class="container">
            <span class="navbar-brand mb-0 h1">
                <i class="fas fa-university"></i>
                CSRR Faculty Tracker
            </span>
            <div class="navbar-nav">
                <span class="nav-link">Rutgers Law School</span>
            </div>
        </div>
    </nav>

    <div class="container mt-4">
        {% with messages = get_flashed_messages(with_categories=true) %}
            {% if messages %}
                {% for category, message in messages %}
                    <div class="alert alert-{% if category == 'error' %}danger{% else %}{{ category }}{% endif %} alert-dismissible fade show">
                        {{ message }}
                        <button type="button" class="btn-close" data-bs-dismiss="alert"></button>
                    </div>
                {% endfor %}
            {% endif %}
        {% endwith %}

        <!-- Hero Section -->
        <div class="search-section">
            <h1><i class="fas fa-search me-3"></i>CSRR Faculty Publication Tracker</h1>
            <p class="lead mb-4">Automated monitoring of publications, op-eds, and media appearances by Center for Security, Race and Rights faculty affiliates</p>
            <form method="POST" action="/run-search" class="d-inline">
                <button type="submit" class="btn btn-light btn-lg px-5">
                    <i class="fas fa-play me-2"></i>Start Monthly Search
                </button>
            </form>
            <p class="mt-3 mb-0"><small>Searches {{ analytics.total_faculty }} faculty members across multiple media sources</small></p>
        </div>

        <!-- Statistics Cards -->
        <div class="row mb-4">
            <div class="col-md-3">
                <div class="card stat-card text-center">
                    <div class="card-body">
                        <i class="fas fa-users fa-3x mb-3"></i>
                        <h3>{{ analytics.total_faculty }}</h3>
                        <p>Faculty Tracked</p>
                    </div>
                </div>
            </div>
            <div class="col-md-3">
                <div class="card stat-card text-center">
                    <div class="card-body">
                        <i class="fas fa-search fa-3x mb-3"></i>
                        <h3>{{ analytics.total_searches }}</h3>
                        <p>Searches Completed</p>
                    </div>
                </div>
            </div>
            <div class="col-md-3">
                <div class="card stat-card text-center">
                    <div class="card-body">
                        <i class="fas fa-envelope fa-3x mb-3"></i>
                        <h3>{{ analytics.total_subscribers }}</h3>
                        <p>Email Subscribers</p>
                    </div>
                </div>
            </div>
            <div class="col-md-3">
                <div class="card stat-card text-center">
                    <div class="card-body">
                        <i class="fas fa-calendar fa-3x mb-3"></i>
                        <h3>{{ analytics.last_search.split()[0] if analytics.last_search != 'Never' else 'Never' }}</h3>
                        <p>Last Search</p>
                    </div>
                </div>
            </div>
        </div>

        <div class="row">
            <!-- Search Results -->
            <div class="col-md-8">
                <div class="card">
                    <div class="card-header bg-white">
                        <h5 class="mb-0"><i class="fas fa-history me-2"></i>Search Results</h5>
                    </div>
                    <div class="card-body">
                        {% if search_history %}
                        <div class="table-responsive">
                            <table class="table table-hover">
                                <thead>
                                    <tr>
                                        <th>Date & Time</th>
                                        <th>Status</th>
                                        <th>Publications Found</th>
                                        <th>Download Reports</th>
                                    </tr>
                                </thead>
                                <tbody>
                                    {% for search in search_history %}
                                    <tr>
                                        <td>{{ search.date }}</td>
                                        <td>
                                            {% if search.status == 'Completed' %}
                                                <span class="badge bg-success">{{ search.status }}</span>
                                            {% elif search.status == 'Running' %}
                                                <span class="badge bg-primary">{{ search.status }}</span>
                                            {% else %}
                                                <span class="badge bg-danger">{{ search.status }}</span>
                                            {% endif %}
                                        </td>
                                        <td>{{ search.results }}</td>
                                        <td>
                                            {% if search.status == 'Completed' %}
                                                <a href="/download-report/{{ search.id }}/excel" class="btn btn-sm btn-outline-success me-1">
                                                    <i class="fas fa-file-excel"></i> Excel
                                                </a>
                                                <a href="/download-report/{{ search.id }}/word" class="btn btn-sm btn-outline-primary">
                                                    <i class="fas fa-file-word"></i> Word
                                                </a>
                                            {% endif %}
                                        </td>
                                    </tr>
                                    {% endfor %}
                                </tbody>
                            </table>
                        </div>
                        {% else %}
                        <div class="text-center py-5">
                            <i class="fas fa-search fa-3x text-muted mb-3"></i>
                            <p class="text-muted">No searches performed yet.</p>
                            <p>Click "Start Monthly Search" above to begin tracking faculty publications!</p>
                        </div>
                        {% endif %}
                    </div>
                </div>
            </div>

            <!-- Email Subscription -->
            <div class="col-md-4">
                <div class="card">
                    <div class="card-header bg-white">
                        <h5 class="mb-0"><i class="fas fa-bell me-2"></i>Monthly Email Reports</h5>
                    </div>
                    <div class="card-body">
                        <p>Receive automated monthly faculty publication reports on the 30th of each month.</p>
                        <form method="POST" action="/subscribe">
                            <div class="input-group mb-3">
                                <input type="email" name="email" class="form-control" placeholder="Enter your email" required>
                                <button type="submit" class="btn btn-primary">
                                    <i class="fas fa-plus"></i>
                                </button>
                            </div>
                        </form>
                        <small class="text-muted">{{ analytics.total_subscribers }} people subscribed</small>
                    </div>
                </div>

                <div class="card">
                    <div class="card-header bg-white">
                        <h5 class="mb-0"><i class="fas fa-users me-2"></i>Faculty Information</h5>
                    </div>
                    <div class="card-body">
                        <p><strong>{{ analytics.total_faculty }}</strong> CSRR faculty affiliates monitored</p>
                        <a href="/faculty-list" class="btn btn-outline-primary btn-sm">
                            <i class="fas fa-list"></i> View Full List
                        </a>
                        <hr>
                        <h6>Content Types Tracked:</h6>
                        <ul class="list-unstyled">
                            <li><i class="fas fa-newspaper text-primary me-2"></i>Op-Eds & Editorials</li>
                            <li><i class="fas fa-microphone text-primary me-2"></i>Interviews</li>
                            <li><i class="fas fa-tv text-primary me-2"></i>TV Appearances</li>
                            <li><i class="fas fa-podcast text-primary me-2"></i>Podcasts</li>
                        </ul>
                    </div>
                </div>
            </div>
        </div>

        <!-- Features -->
        <div class="row mt-5">
            <div class="col-md-4">
                <div class="card text-center h-100">
                    <div class="card-body">
                        <i class="fas fa-robot fa-4x feature-icon mb-3"></i>
                        <h5>Automated Monitoring</h5>
                        <p>Continuous tracking of {{ analytics.total_faculty }} CSRR faculty affiliates across major media outlets and academic publications.</p>
                    </div>
                </div>
            </div>
            <div class="col-md-4">
                <div class="card text-center h-100">
                    <div class="card-body">
                        <i class="fas fa-file-alt fa-4x feature-icon mb-3"></i>
                        <h5>Professional Reports</h5>
                        <p>Generates Excel and Word documents formatted for CSRR team review and website content selection.</p>
                    </div>
                </div>
            </div>
            <div class="col-md-4">
                <div class="card text-center h-100">
                    <div class="card-body">
                        <i class="fas fa-globe fa-4x feature-icon mb-3"></i>
                        <h5>Website Integration</h5>
                        <p>Supports content curation for <a href="https://csrr.rutgers.edu/newsroom/csrr-in-the-news/" target="_blank" style="color: var(--rutgers-red);">CSRR in the News</a> section.</p>
                    </div>
                </div>
            </div>
        </div>
    </div>

    <!-- Footer -->
    <footer class="mt-5 py-4" style="background-color: var(--rutgers-gray); color: white;">
        <div class="container text-center">
            <p class="mb-0">Center for Security, Race and Rights | Rutgers Law School</p>
            <p class="mb-0"><small>Automated Faculty Publication Tracking System</small></p>
        </div>
    </footer>

    <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/js/bootstrap.bundle.min.js"></script>
</body>
</html>
'''

FACULTY_LIST_HTML = '''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>CSRR Faculty Affiliates | Rutgers Law</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/css/bootstrap.min.css" rel="stylesheet">
    <link href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0/css/all.min.css" rel="stylesheet">
    <style>
        :root {
            --rutgers-red: #CC0033;
            --rutgers-light-red: #E74C3C;
            --rutgers-dark-red: #A00025;
        }
        
        .navbar { 
            background: linear-gradient(135deg, var(--rutgers-red) 0%, var(--rutgers-dark-red) 100%);
        }
        
        .card {
            transition: transform 0.2s ease;
        }
        
        .card:hover {
            transform: translateY(-2px);
            box-shadow: 0 4px 15px rgba(204, 0, 51, 0.2);
        }
    </style>
</head>
<body>
    <nav class="navbar navbar-dark">
        <div class="container">
            <a class="navbar-brand" href="/"><i class="fas fa-university"></i> CSRR Faculty Tracker</a>
        </div>
    </nav>

    <div class="container mt-4">
        <div class="row">
            <div class="col-12">
                <h1><i class="fas fa-users"></i> CSRR Faculty Affiliates</h1>
                <p class="lead text-muted">{{ faculty_names|length }} faculty members tracked by the automated monitoring system</p>
            </div>
        </div>
        
        <div class="row">
            {% for name in faculty_names %}
            <div class="col-md-4 col-lg-3 mb-3">
                <div class="card h-100">
                    <div class="card-body">
                        <h6 class="card-title">{{ name }}</h6>
                        <small class="text-muted">Faculty #{{ loop.index }}</small>
                    </div>
                </div>
            </div>
            {% endfor %}
        </div>
        
        <div class="mt-4">
            <a href="/" class="btn btn-primary">
                <i class="fas fa-arrow-left"></i> Back to Dashboard
            </a>
        </div>
    </div>
</body>
</html>
'''

if __name__ == '__main__':
    print("🚀 Starting Rutgers CSRR Faculty Tracker on http://127.0.0.1:3000")
    print("🏛️  Rutgers Law School themed dashboard with real faculty search functionality")
    print("📊 Click 'Start Monthly Search' to run actual searches for all faculty publications")
    app.run(debug=True, port=3000, host='127.0.0.1')
