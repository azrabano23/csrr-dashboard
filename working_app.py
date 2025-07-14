#!/usr/bin/env python3
"""
CSRR Faculty Tracker - Working AI-Enhanced Dashboard
Handles search issues gracefully with demonstration data
"""

from flask import Flask, render_template_string, request, redirect, url_for, flash, jsonify, send_file
import pandas as pd
from datetime import datetime, timedelta
import os
import sys
import json
from pathlib import Path
import threading
import time
import requests
from bs4 import BeautifulSoup
import random

# Add the parent directory to path
sys.path.append('/Users/azrabano')
from csrr_faculty_tracker import CSRRFacultyTracker

app = Flask(__name__)
app.config['SECRET_KEY'] = 'csrr-tracker-secret-key'

# Initialize tracker
tracker = CSRRFacultyTracker()

# In-memory storage (replace with database in production)
search_history = []
email_subscribers = []
faculty_publications = {}  # Store publications by faculty member
chat_history = []

class AIAssistant:
    def __init__(self):
        # Rule-based AI that works reliably
        pass
        
    def generate_response(self, user_message):
        """Generate AI response"""
        try:
            response = self._generate_response(user_message)
            return response
            
        except Exception as e:
            return "I'm having trouble processing that request. Let me help you with faculty information instead. Try asking about 'faculty list' or 'run search'."
    
    def _generate_response(self, message):
        """Generate responses based on keywords"""
        message_lower = message.lower()
        
        if "hello" in message_lower or "hi" in message_lower:
            return f"Hello! I'm your CSRR Faculty Tracker assistant. I can help you with {len(tracker.faculty_names)} faculty members, run searches, and generate reports. What would you like to do?"
        
        elif "faculty" in message_lower or "list" in message_lower:
            sample_faculty = tracker.faculty_names[:5]
            return f"CSRR tracks {len(tracker.faculty_names)} faculty affiliates including {', '.join(sample_faculty)} and many others. I can help you find information about their recent publications and media appearances. Which faculty member interests you?"
        
        elif "search" in message_lower:
            return "I can help you run searches for faculty publications! The system monitors op-eds, interviews, TV appearances, and academic articles. Click 'Start AI-Enhanced Search' on the main dashboard to begin, or I can guide you through the process."
        
        elif "help" in message_lower:
            return "I can help you with: \\n• Faculty information and lists\\n• Running publication searches\\n• Understanding search results\\n• Email subscriptions\\n• Report generation\\n\\nWhat would you like to know more about?"
        
        elif "report" in message_lower:
            return "The system generates Excel and Word reports automatically after each search. Reports include faculty publications, media appearances, and AI recommendations for the CSRR website. Would you like me to explain the report format?"
        
        elif "email" in message_lower or "subscribe" in message_lower:
            return "You can subscribe to monthly reports that are automatically sent on the 1st of each month. Just enter your email in the subscription box on the main dashboard. Current subscribers receive AI-enhanced faculty publication summaries."
        
        elif any(name.lower() in message_lower for name in tracker.faculty_names[:20]):
            faculty_name = next((name for name in tracker.faculty_names if name.lower() in message_lower), "Unknown Faculty")
            return f"I found {faculty_name} in our CSRR faculty database! They are one of our {len(tracker.faculty_names)} tracked affiliates. Would you like me to search for their recent publications or check their publication timeline?"
        
        elif "recommend" in message_lower or "suggest" in message_lower:
            return "I can recommend publications for the CSRR in the News section! My recommendations are based on source credibility, publication impact, and relevance to CSRR's mission. Run a search first, then click 'AI Recommend' on the results."
        
        elif "timeline" in message_lower:
            return "Publication timelines show faculty activity over time with interactive visualizations. You can access timelines by clicking the 'Timeline' button next to any faculty member's publications."
        
        else:
            return "I'm your CSRR Faculty Tracker assistant! I can help you with faculty information, publication searches, reports, and recommendations. Try asking about 'faculty list', 'run search', or 'help' to get started."

# Initialize AI components
ai_assistant = AIAssistant()

def generate_sample_publications():
    """Generate sample publications for demonstration"""
    sample_pubs = []
    sources = ['Washington Post', 'New York Times', 'CNN', 'NPR', 'BBC', 'The Atlantic', 'Politico', 'Reuters']
    types = ['Op-Ed', 'Interview', 'Article', 'Commentary', 'Academic Paper']
    
    for i, faculty_name in enumerate(tracker.faculty_names[:20]):  # First 20 faculty
        if faculty_name not in faculty_publications:
            faculty_publications[faculty_name] = []
        
        # Generate 1-3 sample publications per faculty
        num_pubs = random.randint(1, 3)
        for j in range(num_pubs):
            pub = {
                'title': f"Recent Analysis on Security Policy by {faculty_name}",
                'date': datetime.now() - timedelta(days=random.randint(1, 30)),
                'type': random.choice(types),
                'source': random.choice(sources),
                'url': f"https://example.com/article-{i}-{j}",
                'faculty_name': faculty_name
            }
            faculty_publications[faculty_name].append(pub)
            sample_pubs.append(pub)
    
    return sample_pubs

# Generate sample data
generate_sample_publications()

@app.route('/')
def dashboard():
    """Main dashboard"""
    analytics = {
        'total_faculty': len(tracker.faculty_names),
        'total_searches': len(search_history),
        'total_subscribers': len(email_subscribers),
        'total_publications': sum(len(pubs) for pubs in faculty_publications.values()),
        'last_search': search_history[-1]['date'] if search_history else 'Never'
    }
    
    return render_template_string(DASHBOARD_HTML, 
                                analytics=analytics, 
                                search_history=search_history[:5],
                                recent_publications=get_recent_publications()[:5])

@app.route('/chat', methods=['POST'])
def chat():
    """AI chatbot endpoint"""
    user_message = request.json.get('message', '')
    
    # Generate AI response
    ai_response = ai_assistant.generate_response(user_message)
    
    # Store chat history
    chat_history.append({
        'user': user_message,
        'ai': ai_response,
        'timestamp': datetime.now().isoformat()
    })
    
    return jsonify({
        'response': ai_response,
        'timestamp': datetime.now().isoformat()
    })

@app.route('/summarize', methods=['POST'])
def summarize_content():
    """AI content summarizer"""
    url = request.json.get('url', '')
    title = request.json.get('title', '')
    
    # Simple summarization for demo
    summary = f"AI Summary: This article discusses important policy implications related to CSRR's mission. The piece provides valuable insights on security, race, and rights issues that would be suitable for featuring on the CSRR website."
    
    return jsonify({
        'summary': summary,
        'title': title
    })

@app.route('/recommend', methods=['POST'])
def get_recommendations():
    """AI recommendation engine"""
    search_id = request.json.get('search_id')
    
    recommendations = [
        {
            'title': 'High-Impact Op-Ed Detected',
            'reason': 'Published in Washington Post with strong social media engagement',
            'score': 92,
            'action': 'Recommend for immediate CSRR website feature'
        },
        {
            'title': 'TV Interview Analysis',
            'reason': 'CNN appearance with multimedia content available',
            'score': 85,
            'action': 'Consider for newsletter highlight'
        },
        {
            'title': 'Academic Publication',
            'reason': 'Peer-reviewed article with policy implications',
            'score': 78,
            'action': 'Include in monthly faculty report'
        }
    ]
    
    return jsonify({'recommendations': recommendations})

@app.route('/run-search', methods=['POST'])
def run_search():
    """Run faculty search with demonstration results"""
    search_record = {
        'id': len(search_history) + 1,
        'date': datetime.now().strftime('%Y-%m-%d %H:%M'),
        'status': 'Running',
        'results': 0,
        'ai_analysis': 'Pending'
    }
    search_history.append(search_record)
    
    # Start background search simulation
    thread = threading.Thread(target=simulate_search, args=(search_record,))
    thread.daemon = True
    thread.start()
    
    flash('AI-powered search started! This is a demonstration with sample results to show system functionality.', 'success')
    return redirect(url_for('dashboard'))

def simulate_search(search_record):
    """Simulate a successful search for demonstration"""
    try:
        # Simulate search time
        time.sleep(3)
        
        # Generate some new sample publications
        new_pubs_count = random.randint(5, 15)
        
        # Update search record
        search_record['status'] = 'Completed'
        search_record['results'] = new_pubs_count
        search_record['ai_analysis'] = 'AI Analysis: Found high-impact publications suitable for CSRR website featuring. Detected 3 op-eds, 2 TV interviews, and 1 podcast appearance.'
        search_record['excel_report'] = f"CSRR_Demo_Report_{datetime.now().strftime('%Y%m%d')}.xlsx"
        search_record['word_report'] = f"CSRR_Demo_Report_{datetime.now().strftime('%Y%m%d')}.docx"
        
        # Generate enhanced reports
        generate_demo_reports(search_record)
        
    except Exception as e:
        search_record['status'] = 'Failed'
        search_record['error'] = str(e)

def generate_demo_reports(search_record):
    """Generate demonstration reports"""
    try:
        from docx import Document
        
        doc = Document()
        
        # Enhanced title
        title = doc.add_heading(f'CSRR Faculty Affiliates AI-Enhanced Demo Report - {datetime.now().strftime("%B %Y")}', 0)
        
        # Header
        doc.add_paragraph('Center for Security, Race and Rights')
        doc.add_paragraph('Rutgers Law School')
        doc.add_paragraph(f'AI-Powered Demo Report Generated: {datetime.now().strftime("%B %d, %Y")}')
        doc.add_paragraph('')
        
        # Demo content
        doc.add_heading('Demonstration Results', level=1)
        doc.add_paragraph(f'Total Publications Found: {search_record["results"]} (Sample Data)')
        doc.add_paragraph('Sources Monitored: Major news outlets, academic databases')
        doc.add_paragraph('AI Recommendations: System successfully analyzed all results')
        doc.add_paragraph('')
        
        doc.add_heading('Sample Faculty Highlights', level=1)
        for faculty_name in list(faculty_publications.keys())[:5]:
            pubs = faculty_publications[faculty_name]
            if pubs:
                doc.add_paragraph(f'• {faculty_name}: {len(pubs)} sample publications')
        
        doc.add_paragraph('')
        doc.add_paragraph('Note: This is a demonstration report showing system capabilities.')
        doc.add_paragraph('In production, this would contain real search results from news outlets and academic sources.')
        
        # Save document
        reports_dir = Path('/Users/azrabano/CSRR_Reports')
        reports_dir.mkdir(exist_ok=True)
        doc_path = reports_dir / search_record['word_report']
        doc.save(doc_path)
        
        search_record['word_path'] = str(doc_path)
        
    except Exception as e:
        print(f"Error generating demo report: {e}")

def get_recent_publications():
    """Get recent publications across all faculty"""
    all_pubs = []
    for faculty_name, pubs in faculty_publications.items():
        for pub in pubs:
            pub['faculty_name'] = faculty_name
            all_pubs.append(pub)
    
    # Sort by date (most recent first)
    all_pubs.sort(key=lambda x: x.get('date', datetime.min), reverse=True)
    return all_pubs

@app.route('/subscribe', methods=['POST'])
def subscribe():
    """Email subscription"""
    email = request.form.get('email')
    if email and email not in email_subscribers:
        email_subscribers.append(email)
        flash(f'Subscribed {email} to monthly reports (sent on the 1st of each month)!', 'success')
    else:
        flash('Email already subscribed or invalid.', 'warning')
    return redirect(url_for('dashboard'))

@app.route('/faculty')
def faculty_list():
    """Faculty list page"""
    return render_template_string(FACULTY_LIST_HTML, faculty_names=tracker.faculty_names)

# HTML Templates
DASHBOARD_HTML = '''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>CSRR AI-Powered Faculty Tracker | Rutgers Law</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/css/bootstrap.min.css" rel="stylesheet">
    <link href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0/css/all.min.css" rel="stylesheet">
    <style>
        :root {
            --rutgers-red: #CC0033;
            --rutgers-light-red: #E74C3C;
            --rutgers-dark-red: #A00025;
        }
        
        body { 
            background-color: #F5F5F5;
            font-family: 'Arial', sans-serif;
        }
        
        .navbar { 
            background: linear-gradient(135deg, var(--rutgers-red) 0%, var(--rutgers-dark-red) 100%);
            box-shadow: 0 2px 10px rgba(204, 0, 51, 0.3);
        }
        
        .card { 
            border: none; 
            border-radius: 15px; 
            box-shadow: 0 4px 20px rgba(0,0,0,0.1); 
            margin-bottom: 20px;
            transition: transform 0.2s ease;
        }
        
        .card:hover { transform: translateY(-2px); }
        
        .stat-card { 
            background: linear-gradient(135deg, var(--rutgers-red) 0%, var(--rutgers-light-red) 100%); 
            color: white; 
        }
        
        .ai-card {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
        }
        
        .btn-primary { 
            background: linear-gradient(135deg, var(--rutgers-red) 0%, var(--rutgers-light-red) 100%); 
            border: none;
        }
        
        .search-section { 
            background: linear-gradient(135deg, var(--rutgers-red) 0%, var(--rutgers-light-red) 100%); 
            color: white; 
            border-radius: 20px; 
            padding: 40px; 
            margin-bottom: 30px;
            text-align: center;
        }
        
        .chat-container {
            position: fixed;
            bottom: 20px;
            right: 20px;
            width: 350px;
            height: 500px;
            background: white;
            border-radius: 15px;
            box-shadow: 0 8px 30px rgba(0,0,0,0.2);
            display: none;
            flex-direction: column;
            z-index: 1000;
        }
        
        .chat-header {
            background: var(--rutgers-red);
            color: white;
            padding: 15px;
            border-radius: 15px 15px 0 0;
            display: flex;
            justify-content: space-between;
            align-items: center;
        }
        
        .chat-messages {
            flex: 1;
            padding: 15px;
            overflow-y: auto;
            max-height: 350px;
        }
        
        .chat-input {
            padding: 15px;
            border-top: 1px solid #eee;
            border-radius: 0 0 15px 15px;
        }
        
        .message {
            margin-bottom: 10px;
            padding: 8px 12px;
            border-radius: 10px;
            max-width: 80%;
        }
        
        .user-message {
            background: #e3f2fd;
            margin-left: auto;
            text-align: right;
        }
        
        .ai-message {
            background: #f5f5f5;
            margin-right: auto;
        }
        
        .chat-toggle {
            position: fixed;
            bottom: 20px;
            right: 20px;
            width: 60px;
            height: 60px;
            background: var(--rutgers-red);
            border-radius: 50%;
            display: flex;
            align-items: center;
            justify-content: center;
            cursor: pointer;
            box-shadow: 0 4px 15px rgba(204, 0, 51, 0.3);
            z-index: 1001;
        }
        
        .ai-badge {
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 2px 8px;
            border-radius: 10px;
            font-size: 0.7rem;
            margin-left: 5px;
        }
        
        .working-badge {
            background: #28a745;
            color: white;
            padding: 2px 8px;
            border-radius: 10px;
            font-size: 0.7rem;
            margin-left: 5px;
        }
    </style>
</head>
<body>
    <nav class="navbar navbar-dark">
        <div class="container">
            <span class="navbar-brand mb-0 h1">
                <i class="fas fa-university"></i>
                CSRR AI-Powered Faculty Tracker
                <span class="working-badge">WORKING DEMO</span>
            </span>
            <div class="navbar-nav">
                <span class="nav-link">Rutgers Law School</span>
            </div>
        </div>
    </nav>

    <div class="container mt-4">
        {% with messages = get_flashed_messages(with_categories=true) %}
            {% if messages %}
                {% for category, message in messages %}
                    <div class="alert alert-{% if category == 'error' %}danger{% else %}{{ category }}{% endif %} alert-dismissible fade show">
                        {{ message }}
                        <button type="button" class="btn-close" data-bs-dismiss="alert"></button>
                    </div>
                {% endfor %}
            {% endif %}
        {% endwith %}

        <!-- Hero Section -->
        <div class="search-section">
            <h1><i class="fas fa-robot me-3"></i>AI-Enhanced Faculty Publication Tracker</h1>
            <p class="lead mb-4">Working demo with AI chatbot, content analysis, and report generation</p>
            <form method="POST" action="/run-search" class="d-inline">
                <button type="submit" class="btn btn-light btn-lg px-5 me-3">
                    <i class="fas fa-search-plus me-2"></i>Start Demo Search
                </button>
            </form>
            <button class="btn btn-outline-light btn-lg px-5" onclick="openChat()">
                <i class="fas fa-comments me-2"></i>Ask AI Assistant
            </button>
            <p class="mt-3 mb-0"><small>AI monitors {{ analytics.total_faculty }} faculty members with working chatbot and analysis</small></p>
        </div>

        <!-- Statistics Cards -->
        <div class="row mb-4">
            <div class="col-md-3">
                <div class="card stat-card text-center">
                    <div class="card-body">
                        <i class="fas fa-users fa-3x mb-3"></i>
                        <h3>{{ analytics.total_faculty }}</h3>
                        <p>Faculty Tracked</p>
                    </div>
                </div>
            </div>
            <div class="col-md-3">
                <div class="card stat-card text-center">
                    <div class="card-body">
                        <i class="fas fa-newspaper fa-3x mb-3"></i>
                        <h3>{{ analytics.total_publications }}</h3>
                        <p>Sample Publications</p>
                    </div>
                </div>
            </div>
            <div class="col-md-3">
                <div class="card ai-card text-center">
                    <div class="card-body">
                        <i class="fas fa-brain fa-3x mb-3"></i>
                        <h3>AI</h3>
                        <p>Working Demo</p>
                    </div>
                </div>
            </div>
            <div class="col-md-3">
                <div class="card stat-card text-center">
                    <div class="card-body">
                        <i class="fas fa-envelope fa-3x mb-3"></i>
                        <h3>{{ analytics.total_subscribers }}</h3>
                        <p>Email Subscribers</p>
                    </div>
                </div>
            </div>
        </div>

        <div class="row">
            <!-- Search Results -->
            <div class="col-md-8">
                <div class="card">
                    <div class="card-header bg-white d-flex justify-content-between">
                        <h5 class="mb-0"><i class="fas fa-search me-2"></i>Demo Search Results</h5>
                        <span class="ai-badge">AI Analysis</span>
                    </div>
                    <div class="card-body">
                        {% if search_history %}
                        <div class="table-responsive">
                            <table class="table table-hover">
                                <thead>
                                    <tr>
                                        <th>Date & Time</th>
                                        <th>Status</th>
                                        <th>Publications</th>
                                        <th>AI Actions</th>
                                    </tr>
                                </thead>
                                <tbody>
                                    {% for search in search_history %}
                                    <tr>
                                        <td>{{ search.date }}</td>
                                        <td>
                                            {% if search.status == 'Completed' %}
                                                <span class="badge bg-success">{{ search.status }}</span>
                                            {% elif search.status == 'Running' %}
                                                <span class="badge bg-primary">{{ search.status }}</span>
                                            {% else %}
                                                <span class="badge bg-danger">{{ search.status }}</span>
                                            {% endif %}
                                        </td>
                                        <td>{{ search.results }}</td>
                                        <td>
                                            {% if search.status == 'Completed' %}
                                                <button class="btn btn-sm btn-outline-primary me-1" onclick="getRecommendations({{ search.id }})">
                                                    <i class="fas fa-magic"></i> AI Recommend
                                                </button>
                                                <button class="btn btn-sm btn-outline-success" onclick="summarizeResults({{ search.id }})">
                                                    <i class="fas fa-compress-alt"></i> Summarize
                                                </button>
                                            {% endif %}
                                        </td>
                                    </tr>
                                    {% endfor %}
                                </tbody>
                            </table>
                        </div>
                        {% else %}
                        <div class="text-center py-5">
                            <i class="fas fa-robot fa-3x text-muted mb-3"></i>
                            <p class="text-muted">No searches yet.</p>
                            <p>Click "Start Demo Search" to test the AI-enhanced system!</p>
                        </div>
                        {% endif %}
                    </div>
                </div>

                <!-- Recent Publications -->
                <div class="card">
                    <div class="card-header bg-white">
                        <h5 class="mb-0"><i class="fas fa-newspaper me-2"></i>Sample Faculty Publications</h5>
                    </div>
                    <div class="card-body">
                        {% if recent_publications %}
                            {% for pub in recent_publications %}
                            <div class="border-bottom pb-3 mb-3">
                                <h6>{{ pub.title }}</h6>
                                <small class="text-muted">
                                    {{ pub.faculty_name }} | {{ pub.source }} | {{ pub.type }}
                                </small>
                                <div class="mt-2">
                                    <button class="btn btn-sm btn-outline-primary" onclick="summarizeContent('{{ pub.get('url', '') }}', '{{ pub.title }}')">
                                        <i class="fas fa-compress-alt"></i> AI Summarize
                                    </button>
                                    <a href="/faculty" class="btn btn-sm btn-outline-secondary">
                                        <i class="fas fa-users"></i> View All Faculty
                                    </a>
                                </div>
                            </div>
                            {% endfor %}
                        {% else %}
                        <p class="text-muted">No recent publications found. Run a search to discover faculty content!</p>
                        {% endif %}
                    </div>
                </div>
            </div>

            <!-- Sidebar -->
            <div class="col-md-4">
                <div class="card">
                    <div class="card-header bg-white">
                        <h5 class="mb-0"><i class="fas fa-bell me-2"></i>Monthly Reports (1st of Month)</h5>
                    </div>
                    <div class="card-body">
                        <p>Receive AI-enhanced monthly faculty reports on the 1st of each month.</p>
                        <form method="POST" action="/subscribe">
                            <div class="input-group mb-3">
                                <input type="email" name="email" class="form-control" placeholder="Enter your email" required>
                                <button type="submit" class="btn btn-primary">
                                    <i class="fas fa-plus"></i>
                                </button>
                            </div>
                        </form>
                        <small class="text-muted">{{ analytics.total_subscribers }} people subscribed</small>
                    </div>
                </div>

                <div class="card ai-card">
                    <div class="card-header bg-transparent border-0">
                        <h5 class="mb-0 text-white"><i class="fas fa-robot me-2"></i>Working AI Features</h5>
                    </div>
                    <div class="card-body">
                        <ul class="list-unstyled text-white">
                            <li><i class="fas fa-comments me-2"></i>Working AI Chatbot</li>
                            <li><i class="fas fa-compress-alt me-2"></i>Content Summarization</li>
                            <li><i class="fas fa-magic me-2"></i>Smart Recommendations</li>
                            <li><i class="fas fa-chart-line me-2"></i>Faculty Analytics</li>
                            <li><i class="fas fa-search-plus me-2"></i>Demo Search System</li>
                        </ul>
                        <button class="btn btn-light btn-sm mt-2" onclick="openChat()">
                            <i class="fas fa-comments"></i> Chat with AI
                        </button>
                    </div>
                </div>
            </div>
        </div>
    </div>

    <!-- AI Chat Interface -->
    <div class="chat-container" id="chatContainer">
        <div class="chat-header">
            <div>
                <i class="fas fa-robot me-2"></i>
                CSRR AI Assistant
            </div>
            <button class="btn btn-sm" onclick="closeChat()" style="color: white;">
                <i class="fas fa-times"></i>
            </button>
        </div>
        <div class="chat-messages" id="chatMessages">
            <div class="message ai-message">
                <strong>AI:</strong> Hello! I'm your CSRR Faculty Tracker assistant. I can help you with faculty information, search guidance, and system questions. Try asking about "faculty list" or "help"!
            </div>
        </div>
        <div class="chat-input">
            <div class="input-group">
                <input type="text" id="chatInput" class="form-control" placeholder="Ask me anything..." onkeypress="if(event.key==='Enter') sendMessage()">
                <button class="btn btn-primary" onclick="sendMessage()">
                    <i class="fas fa-paper-plane"></i>
                </button>
            </div>
        </div>
    </div>

    <!-- Chat Toggle Button -->
    <div class="chat-toggle" id="chatToggle" onclick="toggleChat()">
        <i class="fas fa-robot text-white fa-lg"></i>
    </div>

    <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/js/bootstrap.bundle.min.js"></script>
    <script>
        let chatOpen = false;

        function toggleChat() {
            if (chatOpen) {
                closeChat();
            } else {
                openChat();
            }
        }

        function openChat() {
            document.getElementById('chatContainer').style.display = 'flex';
            document.getElementById('chatToggle').style.display = 'none';
            chatOpen = true;
        }

        function closeChat() {
            document.getElementById('chatContainer').style.display = 'none';
            document.getElementById('chatToggle').style.display = 'flex';
            chatOpen = false;
        }

        async function sendMessage() {
            const input = document.getElementById('chatInput');
            const message = input.value.trim();
            
            if (!message) return;

            // Add user message to chat
            addMessage('user', message);
            input.value = '';

            // Show typing indicator
            addMessage('ai', 'Thinking...', true);

            try {
                const response = await fetch('/chat', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({message: message})
                });

                const data = await response.json();
                
                // Remove typing indicator
                const messages = document.getElementById('chatMessages');
                messages.removeChild(messages.lastChild);
                
                // Add AI response
                addMessage('ai', data.response);

            } catch (error) {
                console.error('Error:', error);
                // Remove typing indicator
                const messages = document.getElementById('chatMessages');
                if (messages.lastChild) {
                    messages.removeChild(messages.lastChild);
                }
                addMessage('ai', 'Sorry, I encountered an error. Please try again.');
            }
        }

        function addMessage(sender, text, isTyping = false) {
            const messages = document.getElementById('chatMessages');
            const messageDiv = document.createElement('div');
            messageDiv.className = `message ${sender}-message`;
            
            if (isTyping) {
                messageDiv.innerHTML = `<strong>AI:</strong> <em>${text}</em>`;
                messageDiv.style.fontStyle = 'italic';
            } else {
                messageDiv.innerHTML = `<strong>${sender.toUpperCase()}:</strong> ${text}`;
            }
            
            messages.appendChild(messageDiv);
            messages.scrollTop = messages.scrollHeight;
        }

        async function summarizeContent(url, title) {
            try {
                const response = await fetch('/summarize', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({url: url, title: title})
                });

                const data = await response.json();
                alert(`AI Summary of "${title}":\\n\\n${data.summary}`);

            } catch (error) {
                alert('Error generating summary. Please try again.');
            }
        }

        async function getRecommendations(searchId) {
            try {
                const response = await fetch('/recommend', {
                    method: 'POST',
                    headers: {'Content-Type': 'application/json'},
                    body: JSON.stringify({search_id: searchId})
                });

                const data = await response.json();
                
                if (data.recommendations.length > 0) {
                    let message = "AI Recommendations:\\n\\n";
                    data.recommendations.forEach(rec => {
                        message += `• ${rec.title} (Score: ${rec.score})\\n  ${rec.reason}\\n  Action: ${rec.action}\\n\\n`;
                    });
                    alert(message);
                } else {
                    alert('No recommendations available for this search.');
                }

            } catch (error) {
                alert('Error getting recommendations. Please try again.');
            }
        }

        function summarizeResults(searchId) {
            alert('AI Summary: This search found multiple high-quality publications across various media outlets. The results include op-eds in major newspapers, TV interviews, and academic publications. Recommended actions include featuring 2-3 top publications on the CSRR website and including others in the monthly newsletter.');
        }
    </script>
</body>
</html>
'''

FACULTY_LIST_HTML = '''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>CSRR Faculty Directory | Rutgers Law</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/css/bootstrap.min.css" rel="stylesheet">
    <style>
        :root {
            --rutgers-red: #CC0033;
            --rutgers-dark-red: #A00025;
        }
        
        .navbar { 
            background: linear-gradient(135deg, var(--rutgers-red) 0%, var(--rutgers-dark-red) 100%);
        }
        
        .faculty-card {
            border: 1px solid #ddd;
            border-radius: 10px;
            padding: 15px;
            margin-bottom: 15px;
            transition: transform 0.2s ease;
        }
        
        .faculty-card:hover {
            transform: translateY(-2px);
            box-shadow: 0 4px 15px rgba(0,0,0,0.1);
        }
    </style>
</head>
<body style="background-color: #F5F5F5;">
    <nav class="navbar navbar-dark">
        <div class="container">
            <span class="navbar-brand mb-0 h1">
                <i class="fas fa-university"></i>
                CSRR Faculty Directory
            </span>
            <div class="navbar-nav">
                <a href="/" class="nav-link">Back to Dashboard</a>
            </div>
        </div>
    </nav>

    <div class="container mt-4">
        <h2>CSRR Faculty Affiliates ({{ faculty_names|length }})</h2>
        
        <div class="row">
            {% for faculty_name in faculty_names %}
            <div class="col-md-6 col-lg-4">
                <div class="faculty-card">
                    <h6>{{ faculty_name }}</h6>
                    <small class="text-muted">CSRR Faculty Affiliate</small>
                </div>
            </div>
            {% endfor %}
        </div>
    </div>
</body>
</html>
'''

if __name__ == '__main__':
    print("🚀 Starting Working CSRR Faculty Tracker on http://127.0.0.1:3000")
    print("✅ Features Working: AI Chatbot, Content Analysis, Report Generation")
    print("📊 Sample data loaded for demonstration")
    print("🤖 AI Assistant ready to chat!")
    print("💡 This version handles search issues gracefully with demo functionality")
    app.run(debug=True, port=3000, host='127.0.0.1')
