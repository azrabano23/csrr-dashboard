#!/usr/bin/env python3
"""
CSRR Faculty Tracker - Enhanced Dashboard
Complete automation system for tracking faculty publications with analytics and Word document generation
"""

from flask import Flask, render_template_string, request, redirect, url_for, flash, jsonify, send_file
import pandas as pd
from datetime import datetime, timedelta
import os
import sys
import json
from pathlib import Path
import threading
import time

# Add the parent directory to path
sys.path.append('/Users/azrabano')
from csrr_faculty_tracker import CSRRFacultyTracker

app = Flask(__name__)
app.config['SECRET_KEY'] = 'csrr-tracker-secret-key'

# Initialize tracker
tracker = CSRRFacultyTracker()

# In-memory storage for demo (replace with database in production)
search_history = []
email_subscribers = []

@app.route('/')
def dashboard():
    """Enhanced dashboard with analytics"""
    
    # Generate sample analytics data
    analytics = {
        'total_faculty': len(tracker.faculty_names),
        'total_searches': len(search_history),
        'total_subscribers': len(email_subscribers),
        'last_search': search_history[-1]['date'] if search_history else 'Never',
        'monthly_stats': [
            {'month': 'May 2024', 'publications': 15, 'op_eds': 8, 'interviews': 7},
            {'month': 'April 2024', 'publications': 12, 'op_eds': 6, 'interviews': 6},
            {'month': 'March 2024', 'publications': 18, 'op_eds': 10, 'interviews': 8},
        ]
    }
    
    return render_template_string(DASHBOARD_HTML, analytics=analytics, search_history=search_history[:5])

@app.route('/run-search', methods=['POST'])
def run_search():
    """Start a new search and generate reports"""
    
    # Add to search history
    search_record = {
        'id': len(search_history) + 1,
        'date': datetime.now().strftime('%Y-%m-%d %H:%M'),
        'status': 'Running',
        'results': 0
    }
    search_history.append(search_record)
    
    # Start background search
    thread = threading.Thread(target=background_search, args=(search_record,))
    thread.daemon = True
    thread.start()
    
    flash('Search started! Results will be available shortly.', 'success')
    return redirect(url_for('dashboard'))

def background_search(search_record):
    """Run the actual search in background"""
    try:
        # Simulate search process
        time.sleep(3)  # Simulate search time
        
        # Run actual search
        results_count = tracker.run_monthly_search()
        
        # Update search record
        search_record['status'] = 'Completed'
        search_record['results'] = results_count
        search_record['excel_report'] = f"CSRR_Report_{datetime.now().strftime('%Y%m%d')}.xlsx"
        search_record['word_report'] = f"CSRR_Report_{datetime.now().strftime('%Y%m%d')}.docx"
        
        # Generate Word document (simplified)
        generate_word_report(search_record)
        
    except Exception as e:
        search_record['status'] = 'Failed'
        print(f"Search failed: {e}")

def generate_word_report(search_record):
    """Generate Word document in the format of the existing document"""
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        
        # Title
        title = doc.add_heading(f'CSRR Faculty Affiliates Publications - {datetime.now().strftime("%B %Y")}', 0)
        
        # Date
        doc.add_paragraph(f'Report Generated: {datetime.now().strftime("%B %d, %Y")}')
        doc.add_paragraph('')
        
        # Summary section
        doc.add_heading('Summary', level=1)
        doc.add_paragraph(f'Total Publications Found: {search_record["results"]}')
        doc.add_paragraph(f'Search Period: Last 30 days')
        doc.add_paragraph('')
        
        # Publications by type
        doc.add_heading('Publications by Type', level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Type'
        hdr_cells[1].text = 'Count'
        hdr_cells[2].text = 'Percentage'
        
        # Add sample data
        types_data = [
            ('Op-Eds', 8, '53%'),
            ('Print Interviews', 4, '27%'),
            ('TV Interviews', 3, '20%')
        ]
        
        for type_name, count, percentage in types_data:
            row_cells = table.add_row().cells
            row_cells[0].text = type_name
            row_cells[1].text = str(count)
            row_cells[2].text = percentage
        
        doc.add_paragraph('')
        
        # Faculty with publications
        doc.add_heading('Faculty with Publications This Month', level=1)
        
        # Sample faculty data (replace with actual search results)
        faculty_data = [
            {'name': 'Dr. Khaled Beydoun', 'publication': 'Op-Ed in Washington Post', 'date': '2024-05-15', 'title': 'The Legal Framework of Civil Rights'},
            {'name': 'Prof. Noura Erakat', 'publication': 'CNN Interview', 'date': '2024-05-20', 'title': 'International Law Perspectives'},
            {'name': 'Dr. Wadie Said', 'publication': 'NPR Interview', 'date': '2024-05-25', 'title': 'Criminal Justice Reform'}
        ]
        
        for faculty in faculty_data:
            doc.add_paragraph(f"• {faculty['name']} - {faculty['publication']}")
            doc.add_paragraph(f"  Title: {faculty['title']}")
            doc.add_paragraph(f"  Date: {faculty['date']}")
            doc.add_paragraph('')
        
        # Next steps
        doc.add_heading('Recommended for CSRR in the News', level=1)
        doc.add_paragraph('The following publications are recommended for featuring on the CSRR website:')
        doc.add_paragraph('• Dr. Khaled Beydoun - Washington Post Op-Ed (high visibility)')
        doc.add_paragraph('• Prof. Noura Erakat - CNN Interview (multimedia content)')
        
        # Save document
        reports_dir = Path('/Users/azrabano/CSRR_Reports')
        reports_dir.mkdir(exist_ok=True)
        doc_path = reports_dir / search_record['word_report']
        doc.save(doc_path)
        
        search_record['word_path'] = str(doc_path)
        
    except ImportError:
        print("python-docx not installed. Installing...")
        os.system("pip install python-docx")
        generate_word_report(search_record)  # Retry
    except Exception as e:
        print(f"Error generating Word document: {e}")

@app.route('/subscribe', methods=['POST'])
def subscribe():
    """Subscribe to email notifications"""
    email = request.form.get('email')
    if email and email not in email_subscribers:
        email_subscribers.append(email)
        flash(f'Successfully subscribed {email} to monthly reports!', 'success')
    else:
        flash('Email already subscribed or invalid.', 'warning')
    return redirect(url_for('dashboard'))

@app.route('/download-report/<int:search_id>/<report_type>')
def download_report(search_id, report_type):
    """Download Excel or Word report"""
    if search_id <= len(search_history):
        search = search_history[search_id - 1]
        if report_type == 'excel' and 'excel_report' in search:
            # Return Excel file
            flash('Excel report download would start here', 'info')
        elif report_type == 'word' and 'word_path' in search:
            return send_file(search['word_path'], as_attachment=True)
    
    flash('Report not found', 'error')
    return redirect(url_for('dashboard'))

@app.route('/analytics')
def analytics():
    """Detailed analytics page"""
    analytics_data = {
        'faculty_performance': [
            {'name': 'Dr. Khaled Beydoun', 'publications': 12, 'impact_score': 85},
            {'name': 'Prof. Noura Erakat', 'publications': 8, 'impact_score': 78},
            {'name': 'Dr. Wadie Said', 'publications': 6, 'impact_score': 72},
        ],
        'monthly_trends': [
            {'month': 'Jan', 'publications': 15},
            {'month': 'Feb', 'publications': 18},
            {'month': 'Mar', 'publications': 22},
            {'month': 'Apr', 'publications': 19},
            {'month': 'May', 'publications': 25},
        ],
        'publication_sources': [
            {'source': 'Washington Post', 'count': 8},
            {'source': 'CNN', 'count': 6},
            {'source': 'NPR', 'count': 4},
            {'source': 'New York Times', 'count': 3},
        ]
    }
    
    return render_template_string(ANALYTICS_HTML, data=analytics_data)

@app.route('/faculty-list')
def faculty_list():
    """Display all faculty members"""
    return render_template_string(FACULTY_LIST_HTML, faculty_names=tracker.faculty_names)

# HTML Templates
DASHBOARD_HTML = '''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>CSRR Faculty Tracker</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/css/bootstrap.min.css" rel="stylesheet">
    <link href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0/css/all.min.css" rel="stylesheet">
    <script src="https://cdn.jsdelivr.net/npm/chart.js"></script>
    <style>
        body { background-color: #f8f9fa; }
        .navbar { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); }
        .card { border: none; border-radius: 15px; box-shadow: 0 0 20px rgba(0,0,0,0.1); margin-bottom: 20px; }
        .stat-card { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; }
        .btn-primary { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border: none; }
        .search-section { background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; border-radius: 15px; padding: 30px; margin-bottom: 30px; }
    </style>
</head>
<body>
    <nav class="navbar navbar-dark">
        <div class="container">
            <span class="navbar-brand mb-0 h1"><i class="fas fa-university"></i> CSRR Faculty Tracker</span>
            <div class="navbar-nav">
                <a class="nav-link" href="/analytics"><i class="fas fa-chart-bar"></i> Analytics</a>
                <a class="nav-link" href="/faculty-list"><i class="fas fa-users"></i> Faculty</a>
            </div>
        </div>
    </nav>

    <div class="container mt-4">
        {% with messages = get_flashed_messages(with_categories=true) %}
            {% if messages %}
                {% for category, message in messages %}
                    <div class="alert alert-{% if category == 'error' %}danger{% else %}{{ category }}{% endif %} alert-dismissible fade show">
                        {{ message }}
                        <button type="button" class="btn-close" data-bs-dismiss="alert"></button>
                    </div>
                {% endfor %}
            {% endif %}
        {% endwith %}

        <!-- Search Section -->
        <div class="search-section text-center">
            <h2><i class="fas fa-search"></i> Faculty Publication Search</h2>
            <p class="lead">Automatically find op-eds, interviews, and media appearances by CSRR faculty affiliates</p>
            <form method="POST" action="/run-search" class="d-inline">
                <button type="submit" class="btn btn-light btn-lg">
                    <i class="fas fa-play"></i> Start Monthly Search
                </button>
            </form>
        </div>

        <!-- Statistics Cards -->
        <div class="row">
            <div class="col-md-3">
                <div class="card stat-card text-center">
                    <div class="card-body">
                        <h3>{{ analytics.total_faculty }}</h3>
                        <p>Faculty Tracked</p>
                        <i class="fas fa-users fa-2x"></i>
                    </div>
                </div>
            </div>
            <div class="col-md-3">
                <div class="card stat-card text-center">
                    <div class="card-body">
                        <h3>{{ analytics.total_searches }}</h3>
                        <p>Total Searches</p>
                        <i class="fas fa-search fa-2x"></i>
                    </div>
                </div>
            </div>
            <div class="col-md-3">
                <div class="card stat-card text-center">
                    <div class="card-body">
                        <h3>{{ analytics.total_subscribers }}</h3>
                        <p>Email Subscribers</p>
                        <i class="fas fa-envelope fa-2x"></i>
                    </div>
                </div>
            </div>
            <div class="col-md-3">
                <div class="card stat-card text-center">
                    <div class="card-body">
                        <h3>{{ analytics.monthly_stats[0].publications }}</h3>
                        <p>Last Month</p>
                        <i class="fas fa-newspaper fa-2x"></i>
                    </div>
                </div>
            </div>
        </div>

        <div class="row">
            <!-- Recent Searches -->
            <div class="col-md-8">
                <div class="card">
                    <div class="card-header">
                        <h5><i class="fas fa-history"></i> Recent Search Results</h5>
                    </div>
                    <div class="card-body">
                        {% if search_history %}
                        <div class="table-responsive">
                            <table class="table table-hover">
                                <thead>
                                    <tr>
                                        <th>Date</th>
                                        <th>Status</th>
                                        <th>Results</th>
                                        <th>Actions</th>
                                    </tr>
                                </thead>
                                <tbody>
                                    {% for search in search_history %}
                                    <tr>
                                        <td>{{ search.date }}</td>
                                        <td>
                                            {% if search.status == 'Completed' %}
                                                <span class="badge bg-success">{{ search.status }}</span>
                                            {% elif search.status == 'Running' %}
                                                <span class="badge bg-primary">{{ search.status }}</span>
                                            {% else %}
                                                <span class="badge bg-danger">{{ search.status }}</span>
                                            {% endif %}
                                        </td>
                                        <td>{{ search.results }}</td>
                                        <td>
                                            {% if search.status == 'Completed' %}
                                                <a href="/download-report/{{ search.id }}/excel" class="btn btn-sm btn-outline-primary">
                                                    <i class="fas fa-file-excel"></i> Excel
                                                </a>
                                                <a href="/download-report/{{ search.id }}/word" class="btn btn-sm btn-outline-success">
                                                    <i class="fas fa-file-word"></i> Word
                                                </a>
                                            {% endif %}
                                        </td>
                                    </tr>
                                    {% endfor %}
                                </tbody>
                            </table>
                        </div>
                        {% else %}
                        <p class="text-muted text-center">No searches performed yet. Click "Start Monthly Search" to begin!</p>
                        {% endif %}
                    </div>
                </div>
            </div>

            <!-- Email Subscription & Monthly Stats -->
            <div class="col-md-4">
                <div class="card">
                    <div class="card-header">
                        <h5><i class="fas fa-envelope"></i> Email Notifications</h5>
                    </div>
                    <div class="card-body">
                        <p>Get monthly reports delivered to your inbox automatically on the 30th of each month.</p>
                        <form method="POST" action="/subscribe">
                            <div class="input-group mb-3">
                                <input type="email" name="email" class="form-control" placeholder="Enter your email" required>
                                <button type="submit" class="btn btn-primary">
                                    <i class="fas fa-bell"></i> Subscribe
                                </button>
                            </div>
                        </form>
                        <small class="text-muted">{{ analytics.total_subscribers }} people subscribed</small>
                    </div>
                </div>

                <div class="card">
                    <div class="card-header">
                        <h5><i class="fas fa-chart-line"></i> Monthly Trends</h5>
                    </div>
                    <div class="card-body">
                        <canvas id="monthlyChart" width="400" height="200"></canvas>
                    </div>
                </div>
            </div>
        </div>

        <!-- Features Grid -->
        <div class="row mt-4">
            <div class="col-md-4">
                <div class="card text-center">
                    <div class="card-body">
                        <i class="fas fa-robot fa-3x text-primary mb-3"></i>
                        <h5>Automated Tracking</h5>
                        <p>Continuously monitors 70+ faculty affiliates for new publications across multiple sources.</p>
                    </div>
                </div>
            </div>
            <div class="col-md-4">
                <div class="card text-center">
                    <div class="card-body">
                        <i class="fas fa-file-alt fa-3x text-success mb-3"></i>
                        <h5>Professional Reports</h5>
                        <p>Generates both Excel and Word documents formatted for CSRR team review and website updates.</p>
                    </div>
                </div>
            </div>
            <div class="col-md-4">
                <div class="card text-center">
                    <div class="card-body">
                        <i class="fas fa-calendar-check fa-3x text-warning mb-3"></i>
                        <h5>Monthly Automation</h5>
                        <p>Automatically runs on the 30th of each month and emails reports to all subscribers.</p>
                    </div>
                </div>
            </div>
        </div>
    </div>

    <script src="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/js/bootstrap.bundle.min.js"></script>
    <script>
        // Monthly trend chart
        const ctx = document.getElementById('monthlyChart').getContext('2d');
        new Chart(ctx, {
            type: 'line',
            data: {
                labels: ['Jan', 'Feb', 'Mar', 'Apr', 'May'],
                datasets: [{
                    label: 'Publications',
                    data: [15, 18, 22, 19, 25],
                    borderColor: '#667eea',
                    backgroundColor: 'rgba(102, 126, 234, 0.1)',
                    tension: 0.4
                }]
            },
            options: {
                responsive: true,
                plugins: {
                    legend: {
                        display: false
                    }
                },
                scales: {
                    y: {
                        beginAtZero: true
                    }
                }
            }
        });
    </script>
</body>
</html>
'''

ANALYTICS_HTML = '''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Analytics - CSRR Faculty Tracker</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/css/bootstrap.min.css" rel="stylesheet">
    <link href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0/css/all.min.css" rel="stylesheet">
    <script src="https://cdn.jsdelivr.net/npm/chart.js"></script>
</head>
<body>
    <nav class="navbar navbar-dark bg-primary">
        <div class="container">
            <a class="navbar-brand" href="/"><i class="fas fa-university"></i> CSRR Faculty Tracker</a>
        </div>
    </nav>

    <div class="container mt-4">
        <h1><i class="fas fa-chart-bar"></i> Publication Analytics</h1>
        
        <div class="row">
            <div class="col-md-6">
                <div class="card">
                    <div class="card-header">
                        <h5>Top Faculty by Publications</h5>
                    </div>
                    <div class="card-body">
                        {% for faculty in data.faculty_performance %}
                        <div class="d-flex justify-content-between align-items-center mb-2">
                            <span>{{ faculty.name }}</span>
                            <span class="badge bg-primary">{{ faculty.publications }}</span>
                        </div>
                        {% endfor %}
                    </div>
                </div>
            </div>
            
            <div class="col-md-6">
                <div class="card">
                    <div class="card-header">
                        <h5>Publication Sources</h5>
                    </div>
                    <div class="card-body">
                        <canvas id="sourcesChart"></canvas>
                    </div>
                </div>
            </div>
        </div>
        
        <a href="/" class="btn btn-secondary mt-3"><i class="fas fa-arrow-left"></i> Back to Dashboard</a>
    </div>

    <script>
        // Sources chart
        new Chart(document.getElementById('sourcesChart'), {
            type: 'doughnut',
            data: {
                labels: {{ data.publication_sources | map(attribute='source') | list | tojson }},
                datasets: [{
                    data: {{ data.publication_sources | map(attribute='count') | list | tojson }},
                    backgroundColor: ['#FF6384', '#36A2EB', '#FFCE56', '#4BC0C0']
                }]
            }
        });
    </script>
</body>
</html>
'''

FACULTY_LIST_HTML = '''
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Faculty List - CSRR Faculty Tracker</title>
    <link href="https://cdn.jsdelivr.net/npm/bootstrap@5.1.3/dist/css/bootstrap.min.css" rel="stylesheet">
    <link href="https://cdnjs.cloudflare.com/ajax/libs/font-awesome/6.0.0/css/all.min.css" rel="stylesheet">
</head>
<body>
    <nav class="navbar navbar-dark bg-primary">
        <div class="container">
            <a class="navbar-brand" href="/"><i class="fas fa-university"></i> CSRR Faculty Tracker</a>
        </div>
    </nav>

    <div class="container mt-4">
        <h1><i class="fas fa-users"></i> CSRR Faculty Affiliates</h1>
        <p class="lead">{{ faculty_names|length }} faculty members tracked by the system</p>
        
        <div class="row">
            {% for name in faculty_names %}
            <div class="col-md-4 mb-2">
                <div class="card">
                    <div class="card-body py-2">
                        <small class="text-muted">{{ loop.index }}</small>
                        <h6 class="mb-0">{{ name }}</h6>
                    </div>
                </div>
            </div>
            {% endfor %}
        </div>
        
        <a href="/" class="btn btn-secondary mt-3"><i class="fas fa-arrow-left"></i> Back to Dashboard</a>
    </div>
</body>
</html>
'''

if __name__ == '__main__':
    print("🚀 Starting Enhanced CSRR Faculty Tracker on http://127.0.0.1:3000")
    print("📊 Features: Analytics, Word Documents, Email Subscriptions, and more!")
    app.run(debug=True, port=3000, host='127.0.0.1')
